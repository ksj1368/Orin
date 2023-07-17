from django.shortcuts import render, redirect
from .models import Jasoseo, J_Category, User
from .utils import cat_grammar_final, synonym_word_final, coaching_final, bad_thing_final
from common.decorators import login_message_required
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse, FileResponse
from django.conf import settings
import requests
from django.utils import timezone
import re
from docx import Document
import io
import json
from django.contrib import messages
from django.utils.html import escape
import html

# 메인 교정 페이지 렌더링
@login_message_required
def index(request):
    return render(request, 'correction/index.html')

# 교정 결과 반환
@login_message_required
def correct(request):
    
    context = {
        'jasoseo' : request.POST.get('jasoseo', ''),
        'corrected_content' : '',
        # 'error_message': '',
    }
    
    if request.method == 'POST':
        
        original_content = context['jasoseo']
        corrected_content = original_content    #초기화
        
        if not request.POST.get('option1') and not request.POST.get('option2'):
            # 둘다 선택 되지 않은 경우
            return redirect('correction:index')
        else:
            # 문법교정 선택
            if request.POST.get('option1'):
                corrected_content = cat_grammar_final(corrected_content)
            # 자소서 선택
            if request.POST.get('option2'):
                if not request.POST.get('option1'):
                    corrected_content = escape(corrected_content)
                
                url = f"http://{settings.MODEL_IP}:{settings.MODEL_PORT}/{settings.MODEL_SUBHEAD_URL}"
                data = { 'corrected_content' : corrected_content}
                try:
                    response = requests.post(url, json=data)    # 외부 서버로 요청시에는 보통 requests 라이브러리 많이 쓴다고 한다!
                    response.raise_for_status()  #응답 코드 확인
                    responseData = response.json()
                    subhead_pair = list(zip(responseData['subheadings'],responseData['paragraphs']))
                    corrected_content= ''
                    for subhead,paragraph  in subhead_pair:
                        corrected_content += f'"{subhead}"<br><br>{paragraph}<br><br>'
                    corrected_content = corrected_content.replace('\r\n', '<br>') 
                    corrected_content = corrected_content.replace('<br><br><br>','<br><br>')
                    corrected_content = re.sub('(<br>)+$','',corrected_content) # 교정결과 맨 끝의 연속된 <br>들 다 제거.
                except requests.exceptions.HTTPError as errh:
                    print ("Http Error:",errh)
                    context['error_message'] = 'Http Error : 소제목 생성을 요청하는데 문제가 발생했습니다. 다시 시도해주세요'
                except requests.exceptions.ConnectionError as errc:
                    print ("Error Connecting:",errc)
                    context['error_message'] = 'Error Connecting : 소제목 생성을 요청하는데 문제가 발생했습니다. 다시 시도해주세요'
                except requests.exceptions.Timeout as errt:
                    print ("Timeout Error:",errt)
                    context['error_message'] = 'Timeout Error : 소제목 생성을 요청하는데 문제가 발생했습니다. 다시 시도해주세요'
                except requests.exceptions.RequestException as err:
                    print ("Something went wrong",err)
                    context['error_message'] = '소제목 생성을 요청하는데 문제가 발생했습니다. 다시 시도해주세요.'
            context['corrected_content'] = corrected_content
            # print(corrected_content)
    else:
        return redirect('correction:index')
    
    return render(request, 'correction/correct_result.html',context)

# 카테고리 목록 조회 
@login_required
def categories(request):
    if request.method == 'GET': #카테고리 조회
        categories = J_Category.objects.filter(user = request.user)
        category_names = [category.category_name for category in categories]
        return JsonResponse({'categories':category_names},safe=False)
    
    elif request.method == 'POST': #카테고리 생성
        data = json.loads(request.body)
        category_name = data.get('category_name')
        user = request.user
        
        new_category = J_Category.objects.create(
            category_name = category_name,
            user = user
        )
        return JsonResponse({
            'id': new_category.id,
            'category_name': new_category.category_name,
        }, status=201)  # 201 : 생성됨

# 자소서 보관함에 자기소개서 저장
@login_message_required
def mypage_save(request):
    if request.method == 'POST':
        title = request.POST.get('title')
        content = request.POST.get('content')
        category_name = request.POST.get('category')
        user = request.user
        
        category, _ = J_Category.objects.get_or_create(category_name=category_name, user=user)
        
        jasoseo_length = len(content)
        
        Jasoseo.objects.create(
            title=title,
            content=content,
            jasoseo_length=jasoseo_length,
            category=category,
            user=user
        )
        return JsonResponse({'message': '저장이 완료되었습니다.'})
    
    return JsonResponse({'error': '잘못된 접근 방식입니다.'}, status=400)

# 유의어 조회
def get_synonyms(request):
    word = request.POST.get('word')
    synonyms_list = synonym_word_final(word)  # 유의어 크롤링 함수 호출 
    synonyms_dict = {item[0] : item[1] for item in synonyms_list}   
    return JsonResponse(synonyms_dict)

# 자기소개서 코칭 결과 반환 
def get_coaching(request):
    if request.method == 'POST':
        data = json.loads(request.body) #JSON 데이터를 파이썬 dict형태로 변환
        
        corrected_content = data.get('content') 
        
        sentence_advice = coaching_final(corrected_content) # 문장코칭 (dict)         
        context_advice = bad_thing_final(corrected_content) # 문맥코칭 (String)
        
        good_list = sentence_advice['good']
        bad_list = sentence_advice['bad']
        if context_advice:  # context_advice가 None이 아닌 경우에만 추가
           bad_list.append(context_advice)
        
        feedbacks = []  # {label, 코칭코멘트}
        
        for label, reasons in [('good',good_list), ('bad',bad_list)]:
            for reason in reasons:
                feedbacks.append({'label' : label, 'advice':reason})
    
        return JsonResponse({'feedbacks' : feedbacks})
    return JsonResponse({'error': '잘못된 접근 방식입니다.'}, status=400)

# docs로 파일 저장
@login_message_required
def download_docs(request):
    corrected_content = request.POST.get('text','')
    doc = Document()    
    doc.add_paragraph(corrected_content)
    
    # BytesIO를 사용하여 서버 메모리에 .docx 파일생성(임시저장)
    file = io.BytesIO()
    doc.save(file)
    
    # 파일의 내용을 읽어서 HTTP 응답에 넣어준다
    file.seek(0)   # 주의) 파일 포인터를 다시 파일의 처음으로 이동시켜놔야 처음부터 파일 내용을 읽음
    response = FileResponse(file, as_attachment=True, filename='자기소개서.docx')
    
    return response


def download_docs_get(request, pk):
    corrected_content = Jasoseo.objects.get(pk = pk).content
    doc = Document()    
    doc.add_paragraph(corrected_content)
    
    # BytesIO를 사용하여 서버 메모리에 .docx 파일생성(임시저장)
    file = io.BytesIO()
    doc.save(file)
    
    # 파일의 내용을 읽어서 HTTP 응답에 넣어준다
    file.seek(0)   # 주의) 파일 포인터를 다시 파일의 처음으로 이동시켜놔야 처음부터 파일 내용을 읽음
    response = FileResponse(file, as_attachment=True, filename='자기소개서.docx')    
    return response

@login_required

# 저장한 자기소개서 리스트
def my_manage(request):
    my_jasoseo_list = Jasoseo.objects.filter(user_id = request.user.id).order_by('-created_at')
    my_category_list = J_Category.objects.filter(user_id = request.user.id).distinct()
    context = my_jasoseo_list
    return render(request, 'correction/my_manage.html', {'context':context, 'my_category_list':my_category_list})
    
@login_required
# 저장한 자기소개서 내용
def my_manage_detail(request, pk):
    my_jasoseo_detail = Jasoseo.objects.get(pk = pk)
    my_jasoseo_detail.content = my_jasoseo_detail.content.replace('\r\n', '<br>')
    my_jasoseo_detail.content = my_jasoseo_detail.content.replace('\n', '<br>')
    category_name = J_Category.objects.get(id = my_jasoseo_detail.category_id).category_name
    return render(request, 'correction/my_manage_detail.html', {'my_jasoseo_detail':my_jasoseo_detail, "category_name" : category_name})

# 저장한 자기소개서 삭제
def my_jasoseo_delete(request, pk):
    jasoseo = Jasoseo.objects.get(pk=pk)
    if jasoseo.user_id == request.user.id:
        jasoseo.delete()
        return redirect("/correction/my_manage/")
    else:
        return redirect('/correction/my_manage/{pk}/')
    
# 자기소개서 수정
def my_jasoseo_modify(request, pk):
    jasoseo = Jasoseo.objects.get(pk=pk)
    jasoseo.content = html.unescape(jasoseo.content)
    category_list = J_Category.objects.filter(user_id = request.user.id)
    if request.method == 'POST':
        title = request.POST['title']
        content = request.POST['content']
        category_id = request.POST['category'] 
        jasoseo.title = title
        jasoseo.content = escape(content)
        jasoseo.category_id = J_Category.objects.get(id=category_id)
        jasoseo.save()
    return render(request, 'correction/modify_jasoseo.html', {'jasoseo': jasoseo, 'category_list': category_list})
# 자기소개서 카테고리 삭제
def j_category_delete(request, pk):
    category = J_Category.objects.get(pk=pk)
    if category.user_id == request.user.id:
        category.delete()
        return redirect("/correction/my_manage/")
    else:
        return redirect('/correction/my_manage/')
    
## 자기소개서 생성
@login_required
def create_jasoseo(request):
    category_list = J_Category.objects.filter(user_id=request.user.id)
    if request.method == 'POST':
        category_name = request.POST['category']
        try:
            category = J_Category.objects.filter(category_name=category_name, user_id=request.user.id).first()
            new_jasosoe = Jasoseo.objects.create(
                title=request.POST['title'],
                user_id=request.user.id,
                content=escape(request.POST['content']),
                jasoseo_length=len(request.POST['content']),
                category=category,
                created_at=timezone.now(),
            )
            new_jasosoe.save()
            return redirect('correction:my_manage')
        except J_Category.DoesNotExist:
            # 카테고리가 존재하지 않는 경우에 대한 처리
            pass
    return render(request, 'correction/create_jasoseo.html', {'category_list': category_list})

